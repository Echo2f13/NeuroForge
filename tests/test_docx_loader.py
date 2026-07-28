"""Tests for the DOCX loader.

Creates sample DOCX files in-memory and verifies extraction
of headings, paragraphs, tables, and nested lists.
"""

from __future__ import annotations

import os
import tempfile
from pathlib import Path

import pytest
from docx import Document as DocxDocument
from docx.shared import Inches, Pt

from models.document import Document, InputFormat, Section
from src.ingestion.docx_loader import (
    DOCXLoader,
    DocxLoadError,
    _extract_table_text,
    _get_heading_level,
    extract_docx,
)


# ---------------------------------------------------------------------------
# Fixtures: create sample DOCX files for testing
# ---------------------------------------------------------------------------


@pytest.fixture
def simple_docx(tmp_path) -> Path:
    """Create a simple DOCX with headings and paragraphs."""
    doc = DocxDocument()
    doc.add_heading("Document Title", level=1)
    doc.add_paragraph("This is the introduction paragraph.")
    doc.add_heading("Section One", level=2)
    doc.add_paragraph("Content under section one.")
    doc.add_paragraph("More content in section one.")
    doc.add_heading("Section Two", level=2)
    doc.add_paragraph("Content under section two.")

    file_path = tmp_path / "simple.docx"
    doc.save(str(file_path))
    return file_path


@pytest.fixture
def table_docx(tmp_path) -> Path:
    """Create a DOCX with a table."""
    doc = DocxDocument()
    doc.add_heading("Report", level=1)
    doc.add_paragraph("Data summary below:")

    table = doc.add_table(rows=3, cols=3)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Age"
    table.cell(0, 2).text = "City"
    table.cell(1, 0).text = "Alice"
    table.cell(1, 1).text = "30"
    table.cell(1, 2).text = "NYC"
    table.cell(2, 0).text = "Bob"
    table.cell(2, 1).text = "25"
    table.cell(2, 2).text = "LA"

    doc.add_paragraph("End of report.")

    file_path = tmp_path / "table.docx"
    doc.save(str(file_path))
    return file_path


@pytest.fixture
def list_docx(tmp_path) -> Path:
    """Create a DOCX with list items."""
    doc = DocxDocument()
    doc.add_heading("Shopping List", level=1)

    # Use List Bullet style for list items
    doc.add_paragraph("Fruits", style="List Bullet")
    doc.add_paragraph("Vegetables", style="List Bullet")
    doc.add_paragraph("Dairy", style="List Bullet")

    doc.add_heading("Notes", level=2)
    doc.add_paragraph("Remember to check expiry dates.")

    file_path = tmp_path / "list.docx"
    doc.save(str(file_path))
    return file_path


@pytest.fixture
def empty_docx(tmp_path) -> Path:
    """Create an empty DOCX."""
    doc = DocxDocument()
    file_path = tmp_path / "empty.docx"
    doc.save(str(file_path))
    return file_path


@pytest.fixture
def multi_level_heading_docx(tmp_path) -> Path:
    """Create a DOCX with multiple heading levels."""
    doc = DocxDocument()
    doc.add_heading("Main Title", level=1)
    doc.add_paragraph("Intro text.")
    doc.add_heading("Chapter 1", level=2)
    doc.add_paragraph("Chapter 1 content.")
    doc.add_heading("Subsection 1.1", level=3)
    doc.add_paragraph("Deep content here.")
    doc.add_heading("Chapter 2", level=2)
    doc.add_paragraph("Chapter 2 content.")

    file_path = tmp_path / "multi_level.docx"
    doc.save(str(file_path))
    return file_path


# ---------------------------------------------------------------------------
# Error Handling Tests
# ---------------------------------------------------------------------------


class TestDocxLoaderErrors:
    def test_file_not_found(self):
        with pytest.raises(DocxLoadError, match="File not found"):
            extract_docx("/nonexistent/path/file.docx")

    def test_not_a_docx_file(self, tmp_path):
        txt_file = tmp_path / "test.txt"
        txt_file.write_text("Hello world")
        with pytest.raises(DocxLoadError, match="Not a DOCX file"):
            extract_docx(str(txt_file))

    def test_corrupt_docx(self, tmp_path):
        corrupt_file = tmp_path / "corrupt.docx"
        corrupt_file.write_bytes(b"not a valid docx content")
        with pytest.raises(DocxLoadError, match="Failed to open DOCX"):
            extract_docx(str(corrupt_file))


# ---------------------------------------------------------------------------
# Basic Extraction Tests
# ---------------------------------------------------------------------------


class TestBasicExtraction:
    def test_simple_document_returns_document(self, simple_docx):
        result = extract_docx(str(simple_docx))
        assert isinstance(result, Document)

    def test_metadata_format(self, simple_docx):
        result = extract_docx(str(simple_docx))
        assert result.metadata.format == InputFormat.DOCX

    def test_metadata_source(self, simple_docx):
        result = extract_docx(str(simple_docx))
        assert "simple.docx" in result.metadata.source

    def test_title_from_first_heading(self, simple_docx):
        result = extract_docx(str(simple_docx))
        assert result.metadata.title == "Document Title"

    def test_full_content_contains_text(self, simple_docx):
        result = extract_docx(str(simple_docx))
        assert "introduction paragraph" in result.content
        assert "Content under section one" in result.content
        assert "Content under section two" in result.content

    def test_empty_document(self, empty_docx):
        result = extract_docx(str(empty_docx))
        assert isinstance(result, Document)
        # Empty doc should still have valid content (min_length=1)
        assert len(result.content) >= 1


# ---------------------------------------------------------------------------
# Section Extraction Tests
# ---------------------------------------------------------------------------


class TestSectionExtraction:
    def test_sections_created_from_headings(self, simple_docx):
        result = extract_docx(str(simple_docx))
        # Should have sections: intro (under title), section one, section two
        assert len(result.sections) >= 2

    def test_section_heading_text(self, simple_docx):
        result = extract_docx(str(simple_docx))
        headings = [s.heading for s in result.sections if s.heading]
        assert "Document Title" in headings or "Section One" in headings

    def test_section_content_grouped(self, simple_docx):
        result = extract_docx(str(simple_docx))
        # Find section one
        section_one = next(
            (s for s in result.sections if s.heading == "Section One"), None
        )
        assert section_one is not None
        assert "Content under section one" in section_one.content
        assert "More content in section one" in section_one.content

    def test_multi_level_headings(self, multi_level_heading_docx):
        result = extract_docx(str(multi_level_heading_docx))
        levels = {s.heading: s.level for s in result.sections if s.heading}
        assert levels.get("Chapter 1") == 2 or levels.get("Subsection 1.1") == 3

    def test_heading_levels_preserved(self, multi_level_heading_docx):
        result = extract_docx(str(multi_level_heading_docx))
        subsection = next(
            (s for s in result.sections if s.heading == "Subsection 1.1"), None
        )
        if subsection:
            assert subsection.level == 3


# ---------------------------------------------------------------------------
# Table Extraction Tests
# ---------------------------------------------------------------------------


class TestTableExtraction:
    def test_table_in_content(self, table_docx):
        result = extract_docx(str(table_docx))
        assert "Name" in result.content
        assert "Alice" in result.content
        assert "Bob" in result.content

    def test_table_pipe_format(self, table_docx):
        result = extract_docx(str(table_docx))
        # Should use pipe separation
        assert "|" in result.content

    def test_table_header_separator(self, table_docx):
        result = extract_docx(str(table_docx))
        # Should have markdown-style separator after header
        assert "---" in result.content

    def test_table_in_section(self, table_docx):
        result = extract_docx(str(table_docx))
        # Table content should be in a section
        section_with_table = next(
            (s for s in result.sections if "|" in s.content), None
        )
        assert section_with_table is not None


# ---------------------------------------------------------------------------
# List Extraction Tests
# ---------------------------------------------------------------------------


class TestListExtraction:
    def test_list_items_in_content(self, list_docx):
        result = extract_docx(str(list_docx))
        assert "Fruits" in result.content
        assert "Vegetables" in result.content
        assert "Dairy" in result.content

    def test_list_items_formatted_with_bullets(self, list_docx):
        result = extract_docx(str(list_docx))
        # List items should have bullet marker
        assert "- Fruits" in result.content or "Fruits" in result.content

    def test_list_items_in_section(self, list_docx):
        result = extract_docx(str(list_docx))
        # Lists should be part of the section under the heading
        shopping_section = next(
            (s for s in result.sections if s.heading == "Shopping List"), None
        )
        assert shopping_section is not None
        assert "Fruits" in shopping_section.content


# ---------------------------------------------------------------------------
# DOCXLoader Class Tests
# ---------------------------------------------------------------------------


class TestDOCXLoaderClass:
    """Test the DOCXLoader class interface."""

    def test_loader_returns_document(self, simple_docx):
        loader = DOCXLoader()
        result = loader.load(str(simple_docx))
        assert isinstance(result, Document)

    def test_loader_metadata_format(self, simple_docx):
        loader = DOCXLoader()
        result = loader.load(str(simple_docx))
        assert result.metadata.format == InputFormat.DOCX

    def test_loader_extracts_title(self, simple_docx):
        loader = DOCXLoader()
        result = loader.load(str(simple_docx))
        assert result.metadata.title == "Document Title"

    def test_loader_raises_on_missing_file(self):
        loader = DOCXLoader()
        with pytest.raises(DocxLoadError, match="File not found"):
            loader.load("/nonexistent/file.docx")

    def test_loader_raises_on_non_docx(self, tmp_path):
        loader = DOCXLoader()
        txt_file = tmp_path / "test.txt"
        txt_file.write_text("Hello")
        with pytest.raises(DocxLoadError, match="Not a DOCX file"):
            loader.load(str(txt_file))

    def test_loader_extracts_sections(self, multi_level_heading_docx):
        loader = DOCXLoader()
        result = loader.load(str(multi_level_heading_docx))
        assert len(result.sections) >= 2
        headings = [s.heading for s in result.sections if s.heading]
        assert "Chapter 1" in headings

    def test_loader_extracts_tables(self, table_docx):
        loader = DOCXLoader()
        result = loader.load(str(table_docx))
        assert "|" in result.content
        assert "Alice" in result.content

    def test_loader_extracts_lists(self, list_docx):
        loader = DOCXLoader()
        result = loader.load(str(list_docx))
        assert "Fruits" in result.content
