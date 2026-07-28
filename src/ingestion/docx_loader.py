"""DOCX Loader for NeuroForge.

Extracts structured content from .docx files using python-docx.
Handles headings, paragraphs, tables, and nested lists.

Usage:
    from src.ingestion.docx_loader import DOCXLoader

    loader = DOCXLoader()
    doc = loader.load("path/to/document.docx")

    # Or use the functional API directly:
    from src.ingestion.docx_loader import extract_docx
    doc = extract_docx("path/to/document.docx")
"""

from __future__ import annotations

import os
from pathlib import Path
from typing import Optional

from docx import Document as DocxDocument
from docx.document import Document as DocxDocumentType
from docx.table import Table
from docx.text.paragraph import Paragraph

from models.document import Document, DocumentMetadata, InputFormat, Section


class DocxLoadError(Exception):
    """Raised when a DOCX file cannot be loaded or parsed."""

    pass


def _get_heading_level(paragraph: Paragraph) -> Optional[int]:
    """Extract heading level from a paragraph's style.

    Returns:
        Integer 1-6 for heading styles, None for non-heading paragraphs.
    """
    style_name = paragraph.style.name if paragraph.style else ""
    if style_name.startswith("Heading"):
        try:
            level = int(style_name.replace("Heading", "").strip())
            return min(max(level, 1), 6)
        except ValueError:
            return None
    return None


def _extract_table_text(table: Table) -> str:
    """Convert a DOCX table into pipe-separated structured text.

    Each row becomes a line with cells separated by ' | '.
    A separator line is added after the header row.

    Returns:
        Formatted table string.
    """
    rows = []
    for i, row in enumerate(table.rows):
        cells = [cell.text.strip() for cell in row.cells]
        rows.append("| " + " | ".join(cells) + " |")
        if i == 0:
            # Add separator after header row
            separator = "| " + " | ".join("---" for _ in cells) + " |"
            rows.append(separator)
    return "\n".join(rows)


def _get_list_indent_level(paragraph: Paragraph) -> int:
    """Determine the nesting level of a list item.

    Checks the paragraph's indentation level via the XML numbering
    properties (ilvl). Falls back to left indent measurement.

    Returns:
        Integer indentation level (0-based).
    """
    # Try to get indent level from numbering properties
    pPr = paragraph._p.pPr
    if pPr is not None:
        numPr = pPr.find(
            ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ilvl"
        )
        if numPr is not None:
            val = numPr.get(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val"
            )
            if val is not None:
                try:
                    return int(val)
                except ValueError:
                    pass

    # Fallback: check left indent
    if paragraph.paragraph_format.left_indent:
        # Approximate level from indent (each level ~360000 EMU / ~0.5 inch)
        emu = paragraph.paragraph_format.left_indent
        return max(0, int(emu / 360000))

    return 0


def _is_list_paragraph(paragraph: Paragraph) -> bool:
    """Check if a paragraph is a list item (bulleted or numbered).

    Detects list items by checking for numbering properties in the XML
    or by style name containing 'List'.
    """
    style_name = paragraph.style.name if paragraph.style else ""
    if "List" in style_name:
        return True

    # Check XML numbering properties
    pPr = paragraph._p.pPr
    if pPr is not None:
        numPr = pPr.find(
            ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numId"
        )
        if numPr is not None:
            return True

    return False


def _format_list_item(paragraph: Paragraph) -> str:
    """Format a list item with proper indentation prefix.

    Returns:
        Indented list item string with bullet marker.
    """
    level = _get_list_indent_level(paragraph)
    indent = "  " * level
    text = paragraph.text.strip()
    return f"{indent}- {text}"


def extract_docx(file_path: str) -> Document:
    """Extract structured content from a DOCX file.

    Opens a .docx file and extracts all content including:
    - Paragraphs grouped into sections by headings
    - Tables converted to pipe-separated structured text
    - Lists with proper nesting/indentation

    Args:
        file_path: Path to the .docx file.

    Returns:
        A Document instance with metadata, full content, and sections.

    Raises:
        DocxLoadError: If the file doesn't exist, isn't a .docx, or is corrupt.
    """
    path = Path(file_path)

    # Validate file exists
    if not path.exists():
        raise DocxLoadError(f"File not found: {file_path}")

    # Validate extension
    if path.suffix.lower() != ".docx":
        raise DocxLoadError(f"Not a DOCX file: {file_path}")

    # Try to open the document
    try:
        doc: DocxDocumentType = DocxDocument(str(path))
    except Exception as e:
        raise DocxLoadError(f"Failed to open DOCX file: {file_path}. Error: {e}")

    # Build sections and full content
    sections: list[Section] = []
    full_content_parts: list[str] = []
    title: Optional[str] = None

    # Current section accumulator
    current_heading: Optional[str] = None
    current_level: int = 1
    current_content_parts: list[str] = []

    def _flush_section() -> None:
        """Save the current accumulated section."""
        nonlocal current_heading, current_content_parts
        content = "\n".join(current_content_parts).strip()
        if content:
            sections.append(
                Section(
                    heading=current_heading,
                    content=content,
                    level=current_level,
                )
            )
        current_content_parts.clear()

    # Iterate through document body elements (paragraphs and tables)
    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        if tag == "tbl":
            # Handle table
            table = None
            for t in doc.tables:
                if t._tbl is element:
                    table = t
                    break
            if table:
                table_text = _extract_table_text(table)
                current_content_parts.append(table_text)
                full_content_parts.append(table_text)

        elif tag == "p":
            # Handle paragraph
            paragraph = None
            for p in doc.paragraphs:
                if p._p is element:
                    paragraph = p
                    break

            if paragraph is None:
                continue

            text = paragraph.text.strip()
            heading_level = _get_heading_level(paragraph)

            if heading_level is not None:
                # This is a heading — flush the previous section
                _flush_section()
                current_heading = text
                current_level = heading_level

                # Capture the first heading as document title
                if title is None and text:
                    title = text

                if text:
                    full_content_parts.append(text)

            elif _is_list_paragraph(paragraph) and text:
                # Format as list item
                list_item = _format_list_item(paragraph)
                current_content_parts.append(list_item)
                full_content_parts.append(list_item)

            elif text:
                # Regular paragraph
                current_content_parts.append(text)
                full_content_parts.append(text)

    # Flush the last section
    _flush_section()

    # Build full content
    full_content = "\n".join(full_content_parts)

    # Handle edge case: empty document
    if not full_content.strip():
        full_content = " "  # Document.content requires min_length=1

    # Build metadata
    metadata = DocumentMetadata(
        source=str(path.resolve()),
        format=InputFormat.DOCX,
        title=title,
    )

    # Try to extract author and creation date from core properties
    try:
        core_props = doc.core_properties
        if core_props.author:
            metadata.author = core_props.author
        if core_props.created:
            metadata.created_at = core_props.created.isoformat()
    except Exception:
        pass  # Core properties may not be available

    return Document(
        content=full_content,
        metadata=metadata,
        sections=sections,
    )


class DOCXLoader:
    """Loader class for extracting structured content from DOCX files.

    Provides a class-based interface to the DOCX extraction functionality.
    Extracts text with heading hierarchy, tables as structured text,
    and lists with proper nesting.

    Usage:
        loader = DOCXLoader()
        doc = loader.load("path/to/document.docx")
    """

    def load(self, file_path: str) -> Document:
        """Load and extract content from a DOCX file.

        Opens the specified .docx file and extracts all content including:
        - Paragraph text grouped into sections by heading hierarchy
        - Word heading styles (Heading 1-6) mapped to Section levels
        - Tables converted to pipe-delimited structured text
        - Lists with proper nesting (bullets and numbered)

        Args:
            file_path: Path to the .docx file to load.

        Returns:
            A Document instance with:
            - metadata: source path, format=DOCX, title, author, created_at
            - content: full extracted text
            - sections: structural breakdown by headings

        Raises:
            DocxLoadError: If the file doesn't exist, isn't .docx, or is corrupt.
        """
        return extract_docx(file_path)
